"""Render REPORT.md to REPORT.docx using markdown-it-py + python-docx.

Pure-Python (no pandoc). Handles headings, paragraphs with inline emphasis/code/links,
fenced code blocks, tables, images, ordered/unordered lists, blockquotes, and rules.
Mermaid blocks are rendered as monospace code (image rendering would require a JS runtime).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
SRC_MD = ROOT / "REPORT.md"
DST_DOCX = ROOT / "REPORT.docx"

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
# Page is 8.5" x 11" with 0.6" side margins -> 7.3" usable width.
# Cap image at 6.5" wide and 7.5" tall so it always fits without clipping.
MAX_IMG_WIDTH_IN = 6.5
MAX_IMG_HEIGHT_IN = 7.5


def _add_run(paragraph, text, *, bold=False, italic=False, code=False, link=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = CODE_FONT
        run.font.size = Pt(9)
    if color is not None:
        run.font.color.rgb = color
    return run


def _render_inline(paragraph, tokens, max_image_width_in: float = 5.5):
    """Walk children of an `inline` token and emit runs."""
    bold = italic = code = False
    link = None
    for tok in tokens:
        t = tok.type
        if t == "text":
            _add_run(paragraph, tok.content, bold=bold, italic=italic,
                     color=RGBColor(0x06, 0x5F, 0xD8) if link else None)
        elif t == "strong_open":
            bold = True
        elif t == "strong_close":
            bold = False
        elif t == "em_open":
            italic = True
        elif t == "em_close":
            italic = False
        elif t == "code_inline":
            _add_run(paragraph, tok.content, code=True)
        elif t == "link_open":
            link = tok.attrs.get("href", "")
        elif t == "link_close":
            link = None
        elif t == "softbreak":
            _add_run(paragraph, " ")
        elif t == "hardbreak":
            paragraph.add_run().add_break()
        elif t == "image":
            src = tok.attrs.get("src", "")
            _insert_inline_image(paragraph, src, max_width_in=max_image_width_in)
        elif t == "html_inline":
            # Strip raw HTML tags; keep visible content.
            stripped = re.sub(r"<[^>]+>", "", tok.content)
            if stripped:
                _add_run(paragraph, stripped, bold=bold, italic=italic)


def _scaled_size(img_path: Path):
    """Return (width_in, height_in) bounded by MAX_IMG_WIDTH_IN x MAX_IMG_HEIGHT_IN."""
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    aspect = w_px / h_px if h_px else 1.0
    w_in = MAX_IMG_WIDTH_IN
    h_in = w_in / aspect
    if h_in > MAX_IMG_HEIGHT_IN:
        h_in = MAX_IMG_HEIGHT_IN
        w_in = h_in * aspect
    return w_in, h_in


def _resolve_image(src: str):
    if src.startswith(("http://", "https://")):
        return None
    p = (ROOT / src).resolve()
    return p if p.is_file() else None


def _add_image_block(doc: Document, src: str, alt: str = ""):
    """Insert an image as its own centered paragraph (no inline anchoring)."""
    img_path = _resolve_image(src)
    if img_path is None:
        p = doc.add_paragraph()
        _add_run(p, f"[image: {src}]", italic=True, color=RGBColor(0x88, 0x88, 0x88))
        return
    try:
        w_in, h_in = _scaled_size(img_path)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Inches(w_in), height=Inches(h_in))
        if alt:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(alt)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    except Exception:
        p = doc.add_paragraph()
        _add_run(p, f"[image: {src}]", italic=True, color=RGBColor(0x88, 0x88, 0x88))


def _insert_inline_image(paragraph, src: str, max_width_in: float = 5.5):
    """Fallback for images that appear alongside other inline content."""
    img_path = _resolve_image(src)
    if img_path is not None:
        try:
            w_in, h_in = _scaled_size(img_path)
            cap_w = min(w_in, max_width_in)
            cap_h = h_in * (cap_w / w_in) if w_in else h_in
            paragraph.add_run().add_picture(
                str(img_path), width=Inches(cap_w), height=Inches(cap_h)
            )
            return
        except Exception:
            pass
    _add_run(paragraph, f"[image: {src}]", italic=True, color=RGBColor(0x88, 0x88, 0x88))


def _add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    for line in code.rstrip("\n").split("\n"):
        run = p.add_run(line)
        run.font.name = CODE_FONT
        run.font.size = Pt(9)
        p.add_run().add_break()


def _add_table(doc: Document, header_cells, body_rows):
    cols = max(len(header_cells), max((len(r) for r in body_rows), default=0))
    if cols == 0:
        return
    table = doc.add_table(rows=1 + len(body_rows), cols=cols)
    table.style = "Light Grid Accent 1"
    # Per-column width budget: usable page width / cols, minus a small padding.
    cell_img_max = max(1.5, (MAX_IMG_WIDTH_IN / cols) - 0.2)
    for i, cell_tokens in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        _render_inline(p, cell_tokens, max_image_width_in=cell_img_max)
        for run in p.runs:
            run.bold = True
    for r, row in enumerate(body_rows, start=1):
        for c, cell_tokens in enumerate(row):
            cell = table.rows[r].cells[c]
            p = cell.paragraphs[0]
            _render_inline(p, cell_tokens, max_image_width_in=cell_img_max)
    doc.add_paragraph()


def _setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    for level, size in [(1, 20), (2, 16), (3, 13), (4, 12)]:
        s = styles[f"Heading {level}"]
        s.font.name = BODY_FONT
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    # Widen usable area so 6.5"-wide images fit without clipping.
    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)



def _collect_table(tokens, idx):
    """Starting at a `table_open`, return (header_cells, body_rows, end_idx)."""
    header_cells: list = []
    body_rows: list = []
    current_row: list = []
    in_header = False
    i = idx + 1
    while i < len(tokens) and tokens[i].type != "table_close":
        t = tokens[i].type
        if t == "thead_open":
            in_header = True
        elif t == "thead_close":
            in_header = False
        elif t == "tr_open":
            current_row = []
        elif t == "tr_close":
            if in_header:
                header_cells = current_row
            else:
                body_rows.append(current_row)
        elif t in ("th_open", "td_open"):
            j = i + 1
            while j < len(tokens) and tokens[j].type != "inline":
                j += 1
            cell_children = tokens[j].children if j < len(tokens) else []
            current_row.append(cell_children)
        i += 1
    return header_cells, body_rows, i


def render(md_path: Path = SRC_MD, docx_path: Path = DST_DOCX) -> Path:
    md_text = md_path.read_text(encoding="utf-8")
    md = MarkdownIt("commonmark", {"html": True}).enable("table").enable("strikethrough")
    tokens = md.parse(md_text)

    doc = Document()
    _setup_styles(doc)

    list_stack: list = []  # 'ul' or 'ol'
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        t = tok.type

        if t == "heading_open":
            level = int(tok.tag[1])
            inline = tokens[i + 1]
            p = doc.add_heading(level=min(level, 4))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _render_inline(p, inline.children or [])
            i += 3
            continue

        if t == "paragraph_open":
            inline = tokens[i + 1]
            children = inline.children or []
            # If the paragraph is essentially just image(s) (ignoring whitespace
            # text nodes and softbreaks), emit each image as its own centered
            # block so Word doesn't clip them inside an inline run.
            non_image = [
                c for c in children
                if not (
                    c.type == "image"
                    or (c.type == "text" and not c.content.strip())
                    or c.type in ("softbreak", "hardbreak")
                )
            ]
            images_only = (
                not list_stack
                and not non_image
                and any(c.type == "image" for c in children)
            )
            if images_only:
                for c in children:
                    if c.type == "image":
                        alt = "".join(
                            ch.content for ch in (c.children or []) if ch.type == "text"
                        )
                        _add_image_block(doc, c.attrs.get("src", ""), alt)
                i += 3
                continue
            style = None
            if list_stack:
                style = "List Bullet" if list_stack[-1] == "ul" else "List Number"
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            _render_inline(p, children)
            i += 3
            continue

        if t == "fence" or t == "code_block":
            _add_code_block(doc, tok.content)
            i += 1
            continue

        if t == "hr":
            p = doc.add_paragraph()
            p.add_run("―" * 40).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        if t == "blockquote_open":
            depth = 1
            j = i + 1
            inner_text_runs: list = []
            while j < len(tokens) and depth > 0:
                if tokens[j].type == "blockquote_open":
                    depth += 1
                elif tokens[j].type == "blockquote_close":
                    depth -= 1
                elif tokens[j].type == "inline":
                    inner_text_runs.append(tokens[j].children or [])
                j += 1
            for children in inner_text_runs:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                _render_inline(p, children)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i = j
            continue

        if t == "bullet_list_open":
            list_stack.append("ul")
            i += 1
            continue
        if t == "ordered_list_open":
            list_stack.append("ol")
            i += 1
            continue
        if t in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue
        if t in ("list_item_open", "list_item_close"):
            i += 1
            continue

        if t == "table_open":
            header, body, end = _collect_table(tokens, i)
            _add_table(doc, header, body)
            i = end + 1
            continue

        if t == "html_block":
            i += 1
            continue

        i += 1

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    out = render()
    print(f"Wrote {out} ({out.stat().st_size / 1024:.1f} KB)")
